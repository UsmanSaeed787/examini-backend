import io
from pathlib import Path

from pypdf import PdfReader
from docx import Document

from app.models.material import Material
from app.services.storage_service import StorageService

# Cap per material so a huge upload can't blow up the LLM request (~15k tokens).
MAX_CHARS_PER_MATERIAL = 60_000

# Minimum stripped length to consider extraction successful (catches scanned PDFs).
MIN_EXTRACTED_CHARS = 50

SUPPORTED_TYPES = {"pdf", "docx", "txt"}


def distinct_text_length(text: str) -> int:
    """Length of the text's DISTINCT lines.

    A raw character count cannot tell a real document from a scan: scanner
    apps stamp a watermark on every page, so a 100-page image-only PDF still
    yields hundreds of characters ("CamScanner" x N) and slips past any fixed
    floor. Deduplicating lines first collapses that to one short line while
    leaving genuine prose untouched, so the same floor now measures actual
    content rather than page count.
    """
    seen: set[str] = set()
    for line in text.splitlines():
        stripped = line.strip()
        if stripped:
            seen.add(stripped)
    return sum(len(line) for line in seen)


class ParserService:
    """Extracts plain text from uploaded materials for AI exam generation."""

    @staticmethod
    def extract_text(material: Material) -> str:
        """Fetch the material's file and return its plain text.

        Raises ValueError with a human-readable message on any failure.
        """
        file_type = (material.file_type or Path(material.file_name or "").suffix.lstrip(".")).lower()

        if file_type not in SUPPORTED_TYPES:
            raise ValueError(
                f"File type '{file_type}' of '{material.file_name}' is not supported for "
                f"AI generation (supported: pdf, docx, txt)"
            )

        content = ParserService._fetch_bytes(material)

        if file_type == "pdf":
            text = ParserService._extract_pdf(content, material.file_name)
        elif file_type == "docx":
            text = ParserService._extract_docx(content, material.file_name)
        else:
            text = ParserService._extract_txt(content)

        text = text.strip()
        if distinct_text_length(text) < MIN_EXTRACTED_CHARS:
            raise ValueError(
                f"No readable text found in '{material.file_name}' — it may be a "
                f"scanned/image-only document"
            )

        if len(text) > MAX_CHARS_PER_MATERIAL:
            text = text[:MAX_CHARS_PER_MATERIAL] + "\n[... truncated]"

        return text

    @staticmethod
    def _fetch_bytes(material: Material) -> bytes:
        """Get file bytes from Cloudinary, or legacy local disk for old uploads."""
        if material.file_url.startswith("http"):
            try:
                return StorageService.fetch_file(material.file_url)
            except Exception:
                raise ValueError(f"Could not fetch '{material.file_name}' from storage")

        file_path = Path(material.file_url)
        if not file_path.is_absolute():
            backend_root = Path(__file__).parent.parent.parent
            file_path = backend_root / file_path
        if not file_path.exists():
            raise ValueError(f"File for '{material.file_name}' not found on server")
        return file_path.read_bytes()

    @staticmethod
    def _extract_pdf(content: bytes, file_name: str) -> str:
        try:
            reader = PdfReader(io.BytesIO(content))
            pages = [page.extract_text() or "" for page in reader.pages]
            return "\n".join(pages)
        except Exception:
            raise ValueError(f"Failed to read PDF '{file_name}' — the file may be corrupted or password-protected")

    @staticmethod
    def _extract_docx(content: bytes, file_name: str) -> str:
        try:
            document = Document(io.BytesIO(content))
            parts = [p.text for p in document.paragraphs]
            for table in document.tables:
                for row in table.rows:
                    parts.append(" | ".join(cell.text for cell in row.cells))
            return "\n".join(parts)
        except Exception:
            raise ValueError(f"Failed to read DOCX '{file_name}' — the file may be corrupted")

    @staticmethod
    def _extract_txt(content: bytes) -> str:
        try:
            return content.decode("utf-8")
        except UnicodeDecodeError:
            return content.decode("latin-1")
